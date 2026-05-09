from fpdf import FPDF
from docx import Document

# Content of the Project Abstract
title = "LuminaFinance: Project Abstract"
content = {
    "Overview": (
        "LuminaFinance is a lightweight, high-performance personal finance dashboard "
        "designed to provide users with a comprehensive view of their financial health. "
        "The application allows users to track income and expenses, monitor asset portfolios, "
        "and manage categorized transactions through an intuitive, single-page interface."
    ),
    "Architecture": [
        "The project adopts a monolithic architecture, neatly separating the presentation layer from the business logic while remaining within a single repository for easy deployment and management.",
        "Frontend (Presentation Layer): Built using a vanilla web stack (HTML5, CSS3, JavaScript). It operates as a Single Page Application (SPA), utilizing the native fetch API to communicate asynchronously with the backend. This approach ensures a fast, responsive user experience without the overhead of heavy JavaScript frameworks.",
        "Backend (Application Layer): Powered by Python and the Django framework. It utilizes the Django REST Framework (DRF) to expose robust RESTful API endpoints for the frontend. The backend is responsible for secure user authentication, business logic execution, and serving the static frontend assets.",
        "Database (Persistence Layer): The data layer leverages Django's Object-Relational Mapper (ORM), allowing for flexible database management. It defaults to SQLite for frictionless local development and is configured to seamlessly transition to a MySQL 8 (3NF) relational database for production environments via environment variables."
    ],
    "Key Features": [
        "Interactive Dashboard: A dynamic, vanilla JS-driven interface for visualizing financial data.",
        "Transaction Management: Categorized tracking of expenses and income.",
        "Portfolio Tracking: Management of diverse asset classes (e.g., Cash, Stocks, Crypto, Real Estate).",
        "Environment Adaptability: Ready for both local testing and production deployment out-of-the-box.",
        "Developer-Friendly Seeding: Includes custom management scripts to populate mock data for rapid testing and demonstration."
    ],
    "Conclusion": (
        "LuminaFinance demonstrates a clean, maintainable approach to full-stack web development. "
        "By combining the rapid development capabilities of Django with the lightweight nature "
        "of vanilla frontend technologies, the project delivers a scalable and efficient "
        "platform for personal financial management."
    )
}

# --- Create Word Document (.docx) ---
doc = Document()
doc.add_heading(title, 0)

for section, body in content.items():
    doc.add_heading(section, level=1)
    if isinstance(body, list):
        for item in body:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph(body)

word_file = "LuminaFinance_Abstract.docx"
doc.save(word_file)

# --- Create PDF Document (.pdf) ---
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()
pdf.set_font("Arial", 'B', 16)
pdf.cell(200, 10, txt=title, ln=True, align='C')
pdf.ln(10)

pdf.set_font("Arial", size=12)
for section, body in content.items():
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt=section, ln=True)
    pdf.set_font("Arial", size=11)
    if isinstance(body, list):
        for item in body:
            pdf.multi_cell(0, 7, txt=f"- {item}")
    else:
        pdf.multi_cell(0, 7, txt=body)
    pdf.ln(5)

pdf_file = "LuminaFinance_Abstract.pdf"
pdf.output(pdf_file)

print(f"Files generated: {word_file}, {pdf_file}")